import os
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
from bancodedados.banco_dados import *
from datetime import datetime


class ImprimirFicha():
    def __init__(self, documento, id):
        self.banco = BancoDeDados()
        self.documento = Document(documento)
        self.id = id
        self.nome = self.banco.dql(f"SELECT nome FROM clientes WHERE id = {self.id}")
        self.sobrenome = self.banco.dql(f"SELECT sobrenome FROM clientes WHERE id = {self.id}")
        self.rua = self.banco.dql(f"SELECT rua FROM clientes WHERE id = {self.id}")
        self.numero = self.banco.dql(f"SELECT numero FROM clientes WHERE id = {self.id}")
        self.bairro = self.banco.dql(f"SELECT bairro FROM clientes WHERE id = {self.id}")
        self.cep = self.banco.dql(f"SELECT CEP FROM clientes WHERE id = {self.id}")
        self.cidade = self.banco.dql(f"SELECT cidade FROM clientes WHERE id = {self.id}")
        self.estado = self.banco.dql(f"SELECT estado FROM clientes WHERE id = {self.id}")
        self.cpf = self.banco.dql(f"SELECT cpf FROM clientes WHERE id = {self.id}")
        self.celular = self.banco.dql(f"SELECT celular FROM clientes WHERE id = {self.id}")
        self.obs = self.banco.dql(f"SELECT OBS FROM clientes WHERE id = {self.id}")
        self.imagem = self.banco.dql(f"SELECT imagem FROM clientes WHERE id = {self.id}")
        self.converter_bytes_para_png(self.imagem[0][0])
        
        

    def converter_bytes_para_png(self, imagem_bytes):
        try:
            # Crie um objeto BytesIO a partir dos bytes da imagem
            imagem_io = BytesIO(imagem_bytes)

            # Abra a imagem a partir do objeto BytesIO
            with Image.open(imagem_io) as img:
                # Converte a imagem para o formato PNG
                self.img_png = img.convert("RGBA")
                print("Imagem convertida para PNG com sucesso!")
        except Exception as e:
            print(f"Ocorreu um erro ao converter a imagem: {str(e)}")
            self.img_png = None  

    def substituir_valores(self):
        self.referencias = {
            "AAAA": str(self.nome[0][0]),
            "BBBB": str(self.sobrenome[0][0]),
            "CCCC": str(self.rua[0][0]),
            "DDDD": str(self.numero[0][0]),
            "EEEE": str(self.bairro[0][0]),
            "FFFF": str(self.cep[0][0]),
            "GGGG": str(self.cidade[0][0]),
            "HHHH": str(self.estado[0][0]),
            "IIII": str(self.cpf[0][0]),
            "JJJJ": str(self.celular[0][0]),
            "KKKK": str(self.obs[0][0]),
            "LL": str(datetime.now().day),
            "MM": str(datetime.now().month),
            "NNNN": str(datetime.now().year),
        }
        for paragrafo in self.documento.paragraphs:
            for codigo in self.referencias:
                valor = self.referencias[codigo]                
                paragrafo.text = paragrafo.text.replace(codigo, valor)

        segundo_paragrafo = self.documento.paragraphs[1]
        try:
            if self.img_png:
                imagem_bytesio = BytesIO()
                self.img_png.save(imagem_bytesio, format='PNG')
                run = segundo_paragrafo.add_run()
                run.add_picture(imagem_bytesio, width=Inches(4), height=Inches(3))
                print('imagem inserida com sucesso')
            else:
                print('nenhuma imagem para inserir')
        except Exception as e:
            print(f'ocorreu um erro: {str(e)}')

        self.documento.save(f"clientes/fichas_word/Ficha_{self.referencias['AAAA']}.docx")
        
    


           


